#!/usr/bin/env python3
"""
build_cmb_docx.py
Builds CMB_Residuals_Results_v1.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/parkerlee/Desktop/If Im Right/SCI_Project/meta_pulse/CMB_Residuals_Results_v1.docx"

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size={1:14,2:13,3:12}.get(level,12), bold=True)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(text); set_font(run, size=size, bold=bold)

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin  = sec.bottom_margin = Inches(1.0)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(4)
run = tp.add_run(
    "SCI on CMB Residuals: Searching for Organized Amplitude Modulation\n"
    "in the Planck 2018 Temperature Power Spectrum After ΛCDM Subtraction")
set_font(run, size=14, bold=True)

for line in ["Parker J. Lee", "Independent Researcher · Nashville, TN",
             "US Provisional Patent Application 63/904,444", "2026-05-14"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(line)
    set_font(run, size=11, italic=not line.startswith("2026"))

doc.add_paragraph()

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", 1)
add_para(doc,
    "We applied the Structural Coherence Index (SCI) to the Planck 2018 CMB TT power spectrum "
    "residuals — the difference between the observed D_l spectrum and the best-fit ΛCDM prediction "
    "— to test whether any organized amplitude-envelope modulation persists after standard "
    "cosmological model subtraction. The analysis used the publicly available Planck 2018 data "
    "release (N=2,507 unbinned multipoles, l=2 to 2508) and a 83-bin binned spectrum with "
    "explicit best-fit values. "
    "A critical confound was identified and resolved: raw residuals score CORE (SCI≈0.9955) at "
    "all window sizes because the residual amplitude is modulated by the acoustic peak structure "
    "of the error bars — not by non-ΛCDM physics. "
    "The appropriate analysis uses standardized residuals δD_l/σ_l, which remove this confound. "
    "Standardized residuals score INELIGIBLE at the acoustic (W=75, L=150, z=−0.70) and "
    "broad-scale (W=200, L=400, z=−0.05) windows that are directly relevant to cosmological claims. "
    "A weak positive signal appears at fine scale (W=20, L=40, z=+0.80, SCI=0.672) and is "
    "attributed to Planck's known multipole-to-multipole noise correlations from the data "
    "processing pipeline rather than to non-ΛCDM physics. "
    "The overall result is that the Planck 2018 CMB residuals are consistent with noise across "
    "the scales relevant to the Oscillatory Arrow of Time hypothesis.")

# ── 1. Motivation ─────────────────────────────────────────────────────────────
add_heading(doc, "1. Motivation and Prediction", 1)
add_para(doc,
    "The Oscillatory Arrow of Time proposes that entropy and order alternate as complementary "
    "phases of a universal cycle — that the second law of thermodynamics is local (within an epoch) "
    "rather than global (across the full cycle). At cosmological scales, this predicts that the "
    "envelope of cosmological entropy proxies should show organized amplitude modulation rather "
    "than monotonically increasing entropy with random fluctuations.")
add_para(doc,
    "The CMB temperature power spectrum is the most precisely measured cosmological observable "
    "and provides the most sensitive test of large-scale structure. The key distinction: raw CMB "
    "structure is dominated by known physics (acoustic oscillations, lensing, foreground "
    "contamination, instrument noise). The serious test is on the residuals after standard "
    "cosmological model subtraction — asking whether the remaining structure contains "
    "above-null amplitude-envelope coherence that cannot be attributed to known effects.")
add_para(doc,
    "Null hypothesis: Residuals δD_l = D_l^obs − D_l^ΛCDM are consistent with noise (cosmic "
    "variance + instrument noise). SCI on standardized residuals δD_l/σ_l should score "
    "INELIGIBLE across all window sizes, indicating no organized amplitude modulation beyond "
    "phase-randomized surrogates.")

# ── 2. Data ───────────────────────────────────────────────────────────────────
add_heading(doc, "2. Data and Methods", 1)
add_heading(doc, "2.1 Data Sources", 2)
add_para(doc,
    "All data are from the Planck 2018 public data release R3, accessed through the IRSA "
    "Planck archive (irsa.ipac.caltech.edu/data/Planck/release_3/). No API key or registration "
    "required. Three files were downloaded:")
for item in [
    "COM_PowerSpect_CMB-TT-full_R3.01.txt: unbinned TT power spectrum, l=2..2508 (N=2507). "
    "Columns: l, D_l^obs, −δD_l, +δD_l in units of μK².",
    "COM_PowerSpect_CMB-base-plikHM-TTTEEE-lowl-lowE-lensing-minimum-theory_R3.01.txt: "
    "best-fit ΛCDM model D_l (plikHM+lowl+lowE+lensing), same l grid.",
    "COM_PowerSpect_CMB-TT-binned_R3.01.txt: 83 band-power bins with explicit BestFit column.",
]:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"• {item}")
    set_font(run, size=11)

add_heading(doc, "2.2 Residual Construction", 2)
add_para(doc,
    "Raw residuals: δD_l = D_l^obs − D_l^ΛCDM (μK²). "
    "Standardized residuals: ẑ_l = δD_l / σ_l, where σ_l = mean(−δD_l, +δD_l) is the "
    "symmetric error estimate. Standardized residuals have mean ≈ 0 and standard deviation ≈ 1 "
    "across all l-ranges, confirming proper signal-to-noise weighting.")
add_code(doc,
    "resid_std = (Dl_obs - Dl_theory) / (0.5*(err_lo + err_hi) + 1e-6)\n"
    "# Check: mean=-0.007, std=1.014 across l=2..2508")

add_heading(doc, "2.3 SCI Analysis", 2)
add_para(doc,
    "SCI was applied to four series (unbinned/binned × raw/standardized) at multiple window "
    "sizes. The multipole index l is treated as the 'time' axis: l is a continuous index "
    "from 2 to 2508, and SCI asks whether the Hilbert amplitude envelope of δD_l(l) shows "
    "greater autocorrelation than a phase-randomized surrogate. Phase randomization preserves "
    "the power spectrum of the residual series but destroys any organized envelope structure. "
    "Locked parameters: S=200 surrogates, k=0.9, seed=42.")

# Table: window parameter sets
add_para(doc, "Window parameter sets:", space_before=8, space_after=4)
wt = doc.add_table(rows=4, cols=5)
wt.alignment = WD_TABLE_ALIGNMENT.CENTER; wt.style = "Table Grid"
for j, h in enumerate(["Series", "N", "W", "L", "Rationale"]):
    shade_cell(wt.rows[0].cells[j], "2E4057")
    cell_text(wt.rows[0].cells[j], h, bold=True, size=9)
    wt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
for i, (ser, N, W, L, rat) in enumerate([
    ("Unbinned (fine)",   "2507", "20",  "40",  "Sub-acoustic: l-scale ~20 multipoles"),
    ("Unbinned (acoustic)","2507","75",  "150", "Acoustic peak spacing ~300 → W≈¼ period"),
    ("Unbinned (broad)",  "2507", "200", "400", "Envelope across full spectral range"),
]):
    row = wt.rows[i+1]
    fill = "F8F8F8"
    for j, val in enumerate([ser, N, W, L, rat]):
        shade_cell(row.cells[j], fill)
        cell_text(row.cells[j], val, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT if j in (0,4) else WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Table 1. Window parameters for unbinned analysis. Binned (N=83): W=5,15; L=10,25.",
         space_before=4, space_after=10)

# ── 3. Results ────────────────────────────────────────────────────────────────
add_heading(doc, "3. Results", 1)

add_heading(doc, "3.1 The Confound: Why Raw Residuals Score CORE", 2)
add_para(doc,
    "Raw residuals δD_l (in μK²) score CORE at z=6.0 (clipped) across all window sizes. "
    "This is a confound, not a physics result. The mechanism: Planck's error bars σ_l are "
    "themselves amplitude-modulated by the acoustic peak structure. At the first acoustic peak "
    "(l≈200), σ_l ≈ 432 μK² on average, while between peaks it is much smaller. The acoustic "
    "oscillations in D_l^ΛCDM are partially imprinted in D_l^obs as well, so residuals near "
    "acoustic peaks are systematically larger in absolute value than residuals in acoustic "
    "troughs. The Hilbert envelope of the raw residual series therefore detects the acoustic "
    "periodicity of the error bar structure — not any non-ΛCDM signal.")
add_para(doc,
    "Range-by-range confirmation: the standard deviation of raw residuals in l=30..300 is "
    "470.7 μK² (acoustic peak region, large error bars), in l=300..1000 is 141.0 μK², "
    "and in l=1000..2508 is 51.9 μK² — monotonically decreasing with l, tracking the "
    "declining amplitude of the acoustic oscillations and their associated error bars. "
    "The Hilbert envelope detects this systematic variation as 'coherent amplitude modulation,' "
    "producing a spurious CORE score. All subsequent analysis uses standardized residuals.")

add_heading(doc, "3.2 Primary Results: Standardized Residuals", 2)

# Main results table
rt = doc.add_table(rows=6, cols=7)
rt.alignment = WD_TABLE_ALIGNMENT.CENTER; rt.style = "Table Grid"
for j, h in enumerate(["Series", "W", "L", "c_obs", "gap", "z", "SCI [bucket]"]):
    shade_cell(rt.rows[0].cells[j], "2E4057")
    cell_text(rt.rows[0].cells[j], h, bold=True, size=9)
    rt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

rows_data = [
    ("Unbinned std", "20",  "40",  "0.2666", "+0.0358", "+0.796",  "0.672 [CORE_MID]",   "FFF3CD"),  # yellow - ambiguous
    ("Unbinned std", "75",  "150", "0.1450", "−0.0451", "−0.697",  "0.348 [INELIGIBLE]", "FDECEA"),  # red
    ("Unbinned std", "200", "400", "0.1129", "−0.0029", "−0.049",  "0.489 [INELIGIBLE]", "FDECEA"),
    ("Binned std",   "5",   "10",  "0.1958", "+0.0378", "+0.306",  "0.569 [TACTICAL]",   "FFF3CD"),
    ("Binned std",   "15",  "25",  "0.2360", "+0.1387", "+1.300",  "0.763 [CORE]",        "EAF4EA"),  # green - but caveat
]
for i, (ser, W, L, cobs, gap, z, sci_b, fill) in enumerate(rows_data):
    row = rt.rows[i+1]
    for j, val in enumerate([ser, W, L, cobs, gap, z, sci_b]):
        shade_cell(row.cells[j], fill)
        cell_text(row.cells[j], val, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "Table 2. SCI results for standardized residuals δD_l/σ_l. Yellow = ambiguous. "
    "Red = INELIGIBLE. Green = positive but with caveats (see text).",
    space_before=4, space_after=10)

add_para(doc,
    "The two results most relevant to the Oscillatory Arrow claim are the acoustic-scale "
    "(W=75, L=150) and broad-scale (W=200, L=400) unbinned runs. Both are INELIGIBLE: "
    "z=−0.70 and z=−0.05 respectively. These are the scales at which cosmological organized "
    "modulation would be expected to appear if the Oscillatory Arrow hypothesis is correct. "
    "Their null result is direct evidence against organized amplitude modulation at "
    "cosmologically relevant scales in the CMB residuals.")

add_heading(doc, "3.3 The Fine-Scale Positive (W=20, L=40, z=+0.80)", 2)
add_para(doc,
    "At fine scale (W=20, L=40), standardized residuals show a weak positive SCI=0.672, z=+0.80. "
    "This is above the INELIGIBLE threshold but below CORE (0.75). Three candidate explanations "
    "for this signal, in order of likelihood:")
for i, item in enumerate([
    "Planck noise correlations (most likely): The Planck TT likelihood uses a complex "
    "data processing pipeline (pseudo-C_l estimation, mask power leakage correction, "
    "beam window function deconvolution). Adjacent multipoles are not independent — "
    "the noise covariance matrix has off-diagonal terms at small separations Δl. "
    "These correlations create adjacent-multipole envelope coherence at exactly the "
    "scale W=20, L=40 would detect.",
    "Residual foreground contamination (possible): Point source and dust foreground "
    "subtraction is imperfect. Residual foreground power, if correlated across l at "
    "fine scales, could produce a weak positive SCI signal.",
    "Genuine non-ΛCDM structure (least likely): Modulated primordial oscillations "
    "(e.g., from axion monodromy inflation, cosmic string resonances, or oscillatory "
    "features in the inflaton potential) would appear as fine-scale correlated "
    "residuals. The signal is too weak (z=0.80) and at too fine a scale to distinguish "
    "from the above explanations without a dedicated noise model.",
]):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{i+1}. {item}")
    set_font(run, size=11)

add_heading(doc, "3.4 Binned Analysis Caveat", 2)
add_para(doc,
    "The 83-bin analysis at W=15, L=25 scores CORE (z=1.30, SCI=0.763). This result should "
    "be treated with caution. With N=83 points and a window W=15, each envelope estimate "
    "covers ~18% of the series — a regime where edge effects and the finite-series Hilbert "
    "transform approximation are unreliable. The binned spectrum also substantially compresses "
    "information (83 bins vs. 2507 multipoles) while the binning procedure itself introduces "
    "inter-bin correlations. This result is not used as evidence for or against the hypothesis.")

# ── 4. Discussion ─────────────────────────────────────────────────────────────
add_heading(doc, "4. Discussion", 1)

add_heading(doc, "4.1 Verdict on the Oscillatory Arrow at CMB Scales", 2)
add_para(doc,
    "The standardized residual analysis — the only analysis not subject to the acoustic confound "
    "— shows INELIGIBLE SCI at the acoustic and broad scales that would be relevant to "
    "cosmological organized modulation. This is evidence against the Oscillatory Arrow of Time "
    "hypothesis as a claim about CMB-scale amplitude modulation. The null hypothesis (residuals "
    "are consistent with noise) is not rejected at the scales that matter.")
add_para(doc,
    "This falsification is appropriately scoped. The result does not rule out oscillatory "
    "dynamics at other scales or in other observables. The CMB TT power spectrum is sensitive "
    "to a specific set of physical processes (primordial fluctuations, acoustic oscillations, "
    "reionization, lensing). Other cosmological observables — galaxy clustering entropy at low "
    "redshift, cosmic void evolution, 21-cm fluctuations — could in principle carry different "
    "signals. The CMB test is the most precise available measurement, and its null result "
    "is a significant constraint.")

add_heading(doc, "4.2 What Would Constitute Positive Evidence", 2)
add_para(doc,
    "A convincing positive result would require: (1) Standardized residuals scoring CORE "
    "(z > 2.5, SCI > 0.88) at the acoustic or broad scale (W=75 or W=200), "
    "(2) the same result appearing in multiple independent CMB datasets (WMAP, ACT, SPT), "
    "(3) the result persisting after alternative foreground models and mask choices, and "
    "(4) the signal being stronger in temperature than in E-mode polarization, consistent "
    "with the primordial scalar perturbation channel where the Oscillatory Arrow would "
    "most naturally leave an imprint.")

add_heading(doc, "4.3 The Lab Test Remains Open", 2)
add_para(doc,
    "The Floquet-system bench-top test proposed alongside this analysis remains untested and "
    "is not constrained by the CMB result. If the Oscillatory Arrow operates as a local "
    "information-theoretic or quantum mechanical principle rather than as a cosmological "
    "large-scale structure effect, the CMB would not be the right observable. Floquet matter "
    "experiments — periodically driven quantum systems where entropy proxy time-evolution "
    "can be measured with high precision — represent an independent, cleaner test channel "
    "that is not subject to the noise and foreground complications of CMB analysis.")

# ── 5. Conclusion ─────────────────────────────────────────────────────────────
add_heading(doc, "5. Conclusion", 1)
add_para(doc,
    "SCI applied to Planck 2018 CMB TT power spectrum residuals (δD_l = D_l^obs − D_l^ΛCDM, "
    "standardized by error bars) scores INELIGIBLE at the acoustic scale (W=75, L=150: "
    "z=−0.70) and at the broad scale (W=200, L=400: z=−0.05). A weak fine-scale positive "
    "signal (W=20, L=40: z=+0.80) is most likely attributable to Planck's known noise "
    "correlations from the data reduction pipeline.")
add_para(doc,
    "Raw residuals in μK² score CORE (z=6.0) at all scales, but this is identified as a "
    "confound: the acoustic peak structure of the error bars σ_l imprints amplitude "
    "modulation on the raw residuals that is unrelated to any non-ΛCDM physics. Only "
    "standardized residuals constitute a valid test.")
add_para(doc,
    "The overall result is that the Planck 2018 CMB residuals are consistent with noise "
    "at cosmologically relevant scales. This represents evidence against the Oscillatory "
    "Arrow of Time as a CMB-scale phenomenon, while leaving open the possibility of "
    "the hypothesis operating at other scales, in other observables, or through "
    "laboratory quantum systems.")

# ── Appendix ──────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Appendix: Data and Reproduction", 1)
add_para(doc, "Data (public, no registration required):")
add_code(doc,
    "BASE=https://irsa.ipac.caltech.edu/data/Planck/release_3/ancillary-data/cosmoparams\n"
    "curl $BASE/COM_PowerSpect_CMB-TT-full_R3.01.txt -o planck_TT_full.txt\n"
    "curl $BASE/COM_PowerSpect_CMB-base-plikHM-TTTEEE-lowl-lowE-lensing-minimum-theory_R3.01.txt \\\n"
    "     -o planck_theory.txt\n"
    "curl $BASE/COM_PowerSpect_CMB-TT-binned_R3.01.txt -o planck_TT_binned.txt")
add_para(doc, "Run the analysis:")
add_code(doc,
    "cd '/Users/parkerlee/Desktop/If Im Right/SCI_Project'\n"
    "python3 scripts/sci_cmb_residuals_v1.py\n"
    "# Outputs: results/cmb_residuals_v1/")
add_para(doc,
    "SCI engine: sci_score_v3.py (project root). Packages: numpy, scipy, matplotlib, "
    "astropy 7.2.0, camb, healpy. Python 3.14.")
add_para(doc,
    "Key result files: sci_results.csv (all SCI scores), summary.txt, "
    "plots/cmb_residuals_sci.png, plots/hilbert_envelope.png.")

doc.save(OUT)
print(f"Saved: {OUT}")
