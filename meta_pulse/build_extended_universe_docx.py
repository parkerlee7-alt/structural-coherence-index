#!/usr/bin/env python3
"""build_extended_universe_docx.py — Extended Universe SCI Results v1.docx"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/parkerlee/Desktop/If Im Right/SCI_Project/meta_pulse/Extended_Universe_Results_v1.docx"

def sf(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=6, sa=6, ls=18):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(ls)
    r = p.add_run(text); sf(r); return p

def head(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    sf(r, size={1:14,2:13,3:12}.get(level,12), bold=True)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    r = p.add_run(text); sf(r, name="Courier New", size=9)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def ct(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(text); sf(r, size=size, bold=bold)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin  = sec.bottom_margin = Inches(1.0)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(4)
r = tp.add_run(
    "Does Market Coherence Oscillate? Extended Universe SCI (2010–2026)\n"
    "Testing the ~9-Month Cycle Across Macro Regimes")
sf(r, size=14, bold=True)
for line in ["Parker J. Lee", "Independent Researcher · Nashville, TN",
             "US Provisional Patent Application 63/904,444", "2026-05-15"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line); sf(r, size=11, italic=not line.startswith("2026"))
doc.add_paragraph()

# ── Abstract ──────────────────────────────────────────────────────────────────
head(doc, "Abstract", 1)
para(doc,
    "The SCI-of-SCI experiment (56-month window, Sep 2021 – Apr 2026) found a statistically "
    "significant quasi-periodic oscillation in the universe-wide CORE fraction at lag 9 months "
    "(r=+0.582, p<0.00002). This experiment tests whether that cycle generalizes across the "
    "full available history. We extended the rolling SCI universe to 193 monthly dates "
    "(Apr 2010 – Apr 2026) across 1,339 instruments and recomputed the autocorrelation "
    "function of the CORE fraction at lags 1–24. "
    "The ~9-month cycle is not confirmed in the extended dataset: full-series lag-9 ACF = "
    "+0.071, p=0.34. The regime-level breakdown reveals the cycle was entirely isolated to "
    "the 2023–2026 Current regime (lag-9 r=+0.623, p<0.001, n=40). All five prior macro "
    "regimes show near-zero or negative lag-9 autocorrelation. "
    "The 56-month finding was a period-specific artifact, not a stable market cycle. "
    "The dominant signal in the extended dataset is strong short-term persistence: "
    "lag-1 ACF = +0.672 (p<10⁻⁵), indicating that the market's coherence level trends "
    "rather than oscillates — once the universe enters a structural or noise regime, it "
    "remains there for multiple months before transitioning. The CORE fraction ranged "
    "from 0.14 to 0.99 across the full 16-year history, a far wider excursion than seen "
    "in any single sub-window.")

# ── 1. Background ─────────────────────────────────────────────────────────────
head(doc, "1. Background", 1)
para(doc,
    "The SCI-of-SCI experiment computed the autocorrelation function of the monthly "
    "universe-wide CORE fraction across 56 monthly dates (September 2021 to April 2026). "
    "The CORE fraction is the proportion of the 1,339-instrument universe classified CORE "
    "by the locked SCI pipeline (W=7, L=10, S=40, k=0.9, seed=42) at each monthly "
    "rebalance date. It serves as a single-number proxy for the market's structural "
    "coherence level at each point in time.")
para(doc,
    "The 56-month ACF showed a striking quasi-periodic pattern: lag-1 r=+0.337, "
    "a negative trough at lags 4–5 (r≈−0.35), and a strong positive peak at lag 9 "
    "(r=+0.582, t=4.80, p<0.00002, n=47 pairs). This pattern is consistent with a "
    "damped oscillator with a natural period of approximately 8–10 months.")
para(doc,
    "The prediction to test: if the ~9-month period is a genuine, regime-independent "
    "market cycle, it should appear with r ≥ +0.30 and p < 0.01 in the full 2010–2026 "
    "dataset and should be present in at least 4 of the 6 pre-defined macro regimes.")

# ── 2. Methods ────────────────────────────────────────────────────────────────
head(doc, "2. Methods", 1)
head(doc, "2.1 Data", 2)
para(doc,
    "Daily closing prices for all 1,339 instruments were downloaded from Yahoo Finance "
    "(yfinance) from 2010-01-01 to 2026-05-15. Prices were cached locally for reproducibility. "
    "Monthly rebalance dates were the last business day of each month from April 2010 "
    "through April 2026 (193 dates). At each date, a rolling 90-trading-day window of "
    "returns was used for SCI computation. Tickers without at least 63 trading days of "
    "history at a given date were excluded from that date's universe.")
para(doc,
    "Survivorship bias note: the 1,339 instruments are those present in the 2026 universe. "
    "Pre-2015 computation includes only companies that survived through 2026, biasing "
    "early-period CORE fractions toward established, persistent firms. This is noted but "
    "does not affect the primary goal of testing whether the 9-month cycle replicates.")

head(doc, "2.2 SCI Parameters and Computation", 2)
para(doc,
    "Locked parameters: W=7, L=10, S=40, k=0.9, seed=42 — identical to the Triadic Law "
    "experiment. Computation was parallelized over tickers using fork-context multiprocessing "
    "(8 workers). Each worker computed SCI at all 193 monthly dates for one ticker, "
    "avoiding the need to pickle the full returns dictionary.")
code(doc,
    "# ~262,000 SCI computations total (193 dates × 1,339 tickers)\n"
    "# Parallelized: each worker handles one ticker across all dates\n"
    "r = sci_score_v3(returns_window, w=7, L=10, S=40, k=0.9, seed=42)")

head(doc, "2.3 Regime Definitions", 2)
para(doc,
    "Six macro regimes were defined a priori based on known economic event dates:")

rt = doc.add_table(rows=7, cols=3)
rt.alignment = WD_TABLE_ALIGNMENT.CENTER; rt.style = "Table Grid"
for j, h in enumerate(["Regime", "Dates", "N months"]):
    shade(rt.rows[0].cells[j], "2E4057")
    ct(rt.rows[0].cells[j], h, bold=True, size=10)
    rt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
regime_rows = [
    ("QE / Low-Vol Era",       "2010-01-01 – 2014-12-31", "57"),
    ("Normalization",          "2015-01-01 – 2017-12-31", "36"),
    ("Rate Tightening",        "2018-01-01 – 2019-12-31", "24"),
    ("COVID Shock",            "2020-01-01 – 2020-12-31", "12"),
    ("Post-COVID / Inflation", "2021-01-01 – 2022-12-31", "24"),
    ("Current",                "2023-01-01 – 2026-04-30", "40"),
]
fills = ["F8F8F8", "FAFAFA"] * 3
for i, (name, dates, n) in enumerate(regime_rows):
    row = rt.rows[i+1]
    for j, val in enumerate([name, dates, n]):
        shade(row.cells[j], fills[i])
        ct(row.cells[j], val, size=10,
           align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "Table 1. A priori macro regime definitions.", sb=4, sa=10)

# ── 3. Results ────────────────────────────────────────────────────────────────
head(doc, "3. Results", 1)
head(doc, "3.1 Universe Coverage", 2)
para(doc,
    "193 of 196 monthly dates yielded valid data (≥50 instruments). Mean instruments per "
    "date: 1,196 (lower in early periods as fewer instruments had 90 days of history). "
    "CORE fraction statistics across the full 193-month series: "
    "min=0.139, max=0.987, mean=0.328, SD=0.165. The CORE fraction range is far wider "
    "than in the 56-month sub-window (0.145–0.451), reflecting the large variation in "
    "market coherence across the full post-2010 cycle.")

head(doc, "3.2 Full-Series ACF: The 9-Month Cycle Is Not Present", 2)
para(doc,
    "The full-series ACF (N=193 months) does not show the quasi-periodic 9-month pattern "
    "found in the 56-month window. The dominant signal is strong short-term persistence "
    "at lags 1–2, decaying to near-zero by lag 9:")

# Full ACF table
acf_t = doc.add_table(rows=13, cols=4)
acf_t.alignment = WD_TABLE_ALIGNMENT.CENTER; acf_t.style = "Table Grid"
for j, h in enumerate(["Lag (months)", "ACF (r)", "p-value", "Significant?"]):
    shade(acf_t.rows[0].cells[j], "2E4057")
    ct(acf_t.rows[0].cells[j], h, bold=True, size=9)
    acf_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
acf_data = [
    ("1",  "+0.6723", "<0.00001", "Yes ***", "EAF4EA"),
    ("2",  "+0.2974", "0.00003",  "Yes ***", "EAF4EA"),
    ("3",  "+0.1106", "0.129",    "No",      "F8F8F8"),
    ("4",  "+0.1218", "0.095",    "No",      "F8F8F8"),
    ("5",  "+0.1880", "0.010",    "Yes *",   "EAF4EA"),
    ("6",  "+0.1415", "0.053",    "No",      "F8F8F8"),
    ("7",  "+0.0926", "0.209",    "No",      "F8F8F8"),
    ("8",  "+0.0807", "0.275",    "No",      "F8F8F8"),
    ("9",  "+0.0711", "0.337",    "No",      "FDECEA"),
    ("10", "+0.0442", "0.552",    "No",      "F8F8F8"),
    ("12", "−0.0392", "0.600",    "No",      "F8F8F8"),
    ("18", "+0.0199", "0.794",    "No",      "F8F8F8"),
]
for i, (lag, r, p, sig, fill) in enumerate(acf_data):
    row = acf_t.rows[i+1]
    is_lag9 = lag == "9"
    for j, val in enumerate([lag, r, p, sig]):
        shade(row.cells[j], fill)
        ct(row.cells[j], val, size=10, bold=is_lag9)
para(doc, "Table 2. Full-series ACF of CORE fraction (N=193 months). "
          "Red = lag-9 target (not significant). Green = significant lags.",
     sb=4, sa=10)

head(doc, "3.3 Regime-Level Lag-9 ACF", 2)
para(doc,
    "The regime-level breakdown identifies precisely where the lag-9 signal originates. "
    "It is entirely confined to the 2023–2026 Current regime:")

reg_t = doc.add_table(rows=7, cols=5)
reg_t.alignment = WD_TABLE_ALIGNMENT.CENTER; reg_t.style = "Table Grid"
for j, h in enumerate(["Regime", "N", "Lag-9 r", "p-value", "Significant?"]):
    shade(reg_t.rows[0].cells[j], "2E4057")
    ct(reg_t.rows[0].cells[j], h, bold=True, size=9)
    reg_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
reg_data = [
    ("QE / Low-Vol (2010–2014)",       "57", "−0.182", "0.215", "No",    "FDECEA"),
    ("Normalization (2015–2017)",       "36", "−0.141", "0.483", "No",    "FDECEA"),
    ("Rate Tightening (2018–2019)",     "24", "+0.284", "0.305", "No",    "FFF3CD"),
    ("COVID (2020)",                    "12", "+0.733", "0.476", "No*",   "FFF3CD"),
    ("Post-COVID / Infl. (2021–2022)", "24", "−0.116", "0.682", "No",    "FDECEA"),
    ("Current (2023–2026)",            "40", "+0.623", "0.000", "Yes ***","EAF4EA"),
]
for i, (name, n, r, p, sig, fill) in enumerate(reg_data):
    row = reg_t.rows[i+1]
    for j, val in enumerate([name, n, r, p, sig]):
        shade(row.cells[j], fill)
        ct(row.cells[j], val, size=10,
           align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER,
           bold=(name.startswith("Current")))
para(doc,
    "Table 3. Lag-9 ACF of CORE fraction by macro regime. "
    "* COVID n=12 is too small for reliable lag-9 estimates. "
    "Green = significant. Red = negative or null. Yellow = positive but not significant.",
     sb=4, sa=10)
para(doc,
    "The prediction required lag-9 r ≥ +0.30 with p < 0.01 in the full series AND "
    "in at least 4 of 6 regimes. Neither criterion is met. Only the Current regime "
    "satisfies both conditions. All five prior regimes are null or negative.")

head(doc, "3.4 Why the 56-Month Window Showed r=+0.582", 2)
para(doc,
    "The 56-month window (September 2021 – April 2026) spans exactly the Post-COVID/ "
    "Inflation regime (24 months, lag-9 r=−0.116) and the Current regime (40 months, "
    "lag-9 r=+0.623). The Current regime dominates the combined window because it is "
    "longer, and its strong lag-9 signal pulls the combined ACF upward. When pooled "
    "with the negative Post-COVID period, the net result was r=+0.582 — an artifact "
    "of mixing two regimes with opposite lag-9 signs.")
para(doc,
    "This is a textbook case of a windowing artifact: the ACF at a specific lag can "
    "appear strongly positive in a short window that happens to be dominated by one "
    "regime, then collapse to near-zero when the window is extended to include "
    "regime-diverse history.")

# ── 4. What the Extended Data Does Show ───────────────────────────────────────
head(doc, "4. What the Extended Data Does Show", 1)

head(doc, "4.1 Strong Short-Term Persistence", 2)
para(doc,
    "The dominant signal across all 193 months is not oscillatory but persistent. "
    "Lag-1 ACF = +0.672 (p<10⁻⁵) means that last month's coherence level is the "
    "single best predictor of this month's coherence level. Lag-2 ACF = +0.297 "
    "(p<0.0001) is also significant. By lag 5 the signal has faded to r=+0.188, "
    "and by lag 9 it is indistinguishable from zero.")
para(doc,
    "This is the structure of a mean-reverting trend process, not a fixed-period "
    "oscillator. The market's coherence level has inertia: once it shifts toward "
    "structural or noise regime, it stays there for 1–3 months on average, then "
    "gradually reverts. The half-life of a coherence regime is approximately "
    "2–3 months, consistent with the lag-1 autocorrelation of 0.67.")

head(doc, "4.2 The Full Range of Market Coherence", 2)
para(doc,
    "Extending to 2010 reveals the true amplitude of the coherence cycle. The CORE "
    "fraction reached 0.987 near the peak of the post-crisis recovery (2012–2013), "
    "meaning nearly the entire 1,339-instrument universe was classifying as structurally "
    "coherent. At troughs it fell to 0.139. This ±4× amplitude variation dwarfs "
    "anything visible in the 56-month sub-window (range 0.145–0.451). The universe "
    "undergoes genuine large-scale coherence transitions on multi-year timescales "
    "that shorter windows cannot capture.")
para(doc,
    "Survivorship bias inflates the early-period CORE fraction: the 1,339 instruments "
    "present in 2026 are those that survived and grew, disproportionately the strongest "
    "performers from 2010. These firms likely had above-average structural coherence "
    "throughout the period. The true 2010-era universe-wide CORE fraction is probably "
    "lower than what this analysis shows.")

head(doc, "4.3 The Current Regime Anomaly", 2)
para(doc,
    "The 2023–2026 Current regime does show a genuine, significant lag-9 ACF (r=+0.623, "
    "p<0.001). This is not noise — it is a real pattern specific to this period. "
    "A possible explanation: the post-2023 period saw a specific combination of "
    "structural factors (AI-driven sector rotation, Federal Reserve pause cycles, "
    "index concentration in mega-cap tech) that created a ~9-month rhythmic pattern "
    "in cross-sectional coherence. Whether this pattern continues or dissolves as the "
    "regime changes is an open empirical question that can be monitored in real time "
    "with rolling SCI.")

# ── 5. Conclusion ─────────────────────────────────────────────────────────────
head(doc, "5. Conclusion", 1)
para(doc,
    "The ~9-month quasi-periodic oscillation in the universe-wide CORE fraction, "
    "found in the 56-month (2021–2026) window, does not replicate in the extended "
    "2010–2026 dataset. Full-series lag-9 ACF = +0.071, p=0.34 (N=193 months). "
    "The cycle is absent in five of six pre-defined macro regimes and is significant "
    "only in the 2023–2026 Current regime (r=+0.623, p<0.001). The 56-month finding "
    "was a windowing artifact driven by that regime's dominance of the shorter window.")
para(doc,
    "The extended dataset reveals a different and more robust structure: strong "
    "short-term persistence (lag-1 r=+0.672). The market's coherence level trends "
    "rather than oscillates at a fixed period. It has inertia — regimes persist for "
    "months — but no fixed-period cycle structure visible across the full 16-year history.")
para(doc,
    "The full-history CORE fraction range (0.14–0.99) demonstrates that the universe "
    "undergoes genuine large-amplitude coherence transitions over multi-year timescales. "
    "These transitions are the macroscopic signature of what SCI detects at the "
    "instrument level: the market periodically enters and exits structural coherence "
    "at a scale visible across the entire cross-section, but the timing of those "
    "transitions is not governed by a fixed oscillatory clock.")

# ── Appendix ──────────────────────────────────────────────────────────────────
doc.add_page_break()
head(doc, "Appendix: Reproduction", 1)
code(doc,
    "cd '/Users/parkerlee/Desktop/If Im Right/SCI_Project'\n"
    "python3 scripts/finance/sci_universe_extended_v1.py\n"
    "# Downloads 1,339 tickers from 2010-01-01 via yfinance (cached after first run)\n"
    "# Outputs: results/sci_universe_extended_v1/")
para(doc, "Output files:")
code(doc,
    "bucket_fractions_extended.csv   # 193-month CORE/TACTICAL/INELIGIBLE fractions\n"
    "regime_lag9_acf.csv             # lag-9 ACF per macro regime\n"
    "summary.txt                     # full ACF table and summary\n"
    "plots/core_fraction_acf_extended.png")
para(doc,
    "SCI parameters: W=7, L=10, S=40, k=0.9, seed=42. "
    "Parallelization: fork-context multiprocessing, 8 workers, over tickers. "
    "Survivorship bias: 1,339 instruments are those present in the 2026 universe.")

doc.save(OUT)
print(f"Saved: {OUT}")
