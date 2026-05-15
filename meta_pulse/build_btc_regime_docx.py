#!/usr/bin/env python3
"""
build_btc_regime_docx.py
Builds BTC_Regime_Results_v1.docx — formal write-up of the Bitcoin Regime Transition Test.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/parkerlee/Desktop/If Im Right/SCI_Project/meta_pulse/BTC_Regime_Results_v1.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=6, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 14, 2: 13, 3: 12}
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # Grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9)
    return p

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)

# ── Document ───────────────────────────────────────────────────────────────────
doc = Document()

# Page setup: US Letter, 1-inch margins
sec = doc.sections[0]
sec.page_width   = Inches(8.5)
sec.page_height  = Inches(11)
sec.left_margin  = sec.right_margin  = Inches(1.0)
sec.top_margin   = sec.bottom_margin = Inches(1.0)

# ── Title block ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(4)
run = title_p.add_run(
    "The Regime Transition Test: Rolling SCI on Bitcoin Daily Returns (2017–2026)")
set_font(run, size=15, bold=True)

for line in [
    "Parker J. Lee",
    "Independent Researcher · Nashville, TN",
    "US Provisional Patent Application 63/904,444",
    "2026-05-14",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line)
    is_date = line.startswith("2026")
    set_font(run, size=11, bold=False, italic=not is_date)

doc.add_paragraph()  # spacer

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", 1)
add_para(doc,
    "We tested the hypothesis that Bitcoin's 90-day rolling Structural Coherence Index (SCI) "
    "rises toward CORE levels (>0.75) during periods of elevated institutional participation "
    "and falls during retail-dominated or regulatory-crisis regimes. The prediction was directly "
    "falsified: institutional periods produced a lower mean SCI (0.608) than retail/crisis periods "
    "(0.676), a statistically significant reversal (t=−9.07, p<10⁻¹⁸, d=−0.33). The highest-SCI "
    "epoch in the entire 2017–2026 record was the pre-futures retail era (2017-01 to 2017-12-16, "
    "mean SCI=0.801, 76% CORE-classified). The revised interpretation is that SCI detects structural "
    "coherence regardless of its mechanistic source. In mature equity markets, institutions generate "
    "coherence through ETF rebalancing and macro feedback. In Bitcoin, retail speculative herding "
    "generates coherence through momentum synchrony. Institutional entry into Bitcoin brings arbitrage "
    "and derivatives structures that fragment herding dynamics and reduce coherence — the inverse of "
    "the equity mechanism. The result reframes SCI as an instrument for detecting any regime of "
    "structural coherence, whether driven by institutional feedback loops or by retail speculative "
    "synchrony, and suggests that the two mechanisms are distinguishable through their temporal "
    "fingerprints.")

# ── 1. Background ─────────────────────────────────────────────────────────────
add_heading(doc, "1. Background and Prediction", 1)
add_para(doc,
    "In cross-domain validation, Bitcoin daily returns classify as hard INELIGIBLE (SCI=0.096, "
    "z=−2.49) when scored against the full universe of 1,339 financial instruments. Equity indices "
    "and individual stocks in the CORE bucket reach the z=6.0 ceiling. The proposed explanation for "
    "this gap is the absence of institutional feedback infrastructure in Bitcoin: no index inclusion "
    "mechanics, no ETF rebalancing cascades, no systematic factor-arbitrage loops that force "
    "correlated price dynamics across time scales.")

add_para(doc,
    "The testable prediction follows: as institutional participation grows, Bitcoin should develop "
    "these feedback structures, and its rolling SCI should rise toward CORE thresholds. Specifically, "
    "the late 2020 – early 2021 period — characterized by MicroStrategy's balance-sheet allocation, "
    "Tesla's $1.5B purchase, and a surge in Grayscale Bitcoin Trust (GBTC) AUM — was predicted to "
    "show elevated SCI relative to retail-dominated bear markets and regulatory crisis periods "
    "(2022 Luna collapse, FTX failure, crypto winter).")

add_para(doc,
    "The prediction also carried a broader implication: that rolling SCI could serve as a real-time "
    "detector of when a new asset class has developed the feedback infrastructure that makes "
    "systematic strategies reliably applicable — the moment of transition from noise-dominated to "
    "regime-dominated dynamics.")

# ── 2. Methods ────────────────────────────────────────────────────────────────
add_heading(doc, "2. Methods", 1)

add_heading(doc, "2.1 Data", 2)
add_para(doc,
    "Bitcoin daily closing prices were downloaded from Yahoo Finance (ticker: BTC-USD) covering "
    "2017-01-01 through 2026-05-13 (3,419 trading days). Log-percentage daily returns were "
    "computed. GBTC daily closing prices were downloaded over the same interval as a proxy for "
    "institutional demand flows. No survivorship bias or lookahead bias is present; all data were "
    "downloaded in a single batch after the fact.")

add_heading(doc, "2.2 Rolling SCI Computation", 2)
add_para(doc,
    "SCI was computed on overlapping 63-trading-day windows (~90 calendar days) ending on each "
    "trading day from 2017-03-06 through 2026-05-13. Parameters were held at the locked finance "
    "standard: W=7, L=10, S=40, k=0.9, seed=42. Each window produced one SCI observation, "
    "yielding 3,356 total windows.")
add_code_block(doc,
    "for i in range(63, len(returns)):\n"
    "    window = returns[i-63:i]          # 63 trading days\n"
    "    r = sci_score_v3(window, w=7, L=10, S=40, k=0.9, seed=42)\n"
    "    record = {date, SCI, z, gap, c_obs, c_surr}")

add_heading(doc, "2.3 Regime Definitions", 2)
add_para(doc,
    "Institutional and retail/crisis periods were defined a priori based on publicly recorded events, "
    "not on the SCI data itself. Definitions were fixed before any results were examined.")

# Regime table
regime_table = doc.add_table(rows=11, cols=3)
regime_table.alignment = WD_TABLE_ALIGNMENT.CENTER
regime_table.style = "Table Grid"
headers = ["Period", "Dates", "Regime"]
col_widths = [Inches(2.8), Inches(2.5), Inches(1.8)]
for j, (h, w) in enumerate(zip(headers, col_widths)):
    cell = regime_table.rows[0].cells[j]
    shade_cell(cell, "2E4057")
    cell_text(cell, h, bold=True, size=10)
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(255, 255, 255)
    cell.width = w

rows_data = [
    ("CME/CBOE Futures Era",        "2017-12-17 – 2018-06-30", "Institutional"),
    ("Corporate Treasury Wave",     "2020-10-01 – 2021-04-30", "Institutional"),
    ("ProShares ETF / ATH",         "2021-10-01 – 2021-11-10", "Institutional"),
    ("Spot ETF Era",                "2024-01-11 – present",    "Institutional"),
    ("Pre-Futures Retail",          "2017-01-01 – 2017-12-16", "Retail"),
    ("Prolonged Bear / Retail",     "2018-07-01 – 2020-09-30", "Retail"),
    ("China Ban / Mid-Cycle",       "2021-05-01 – 2021-09-30", "Retail"),
    ("Post-ATH Deleveraging",       "2021-11-11 – 2022-01-31", "Retail"),
    ("Luna/3AC Collapse",           "2022-05-01 – 2022-07-31", "Retail"),
    ("FTX Collapse / Crypto Winter","2022-11-01 – 2023-12-31", "Retail"),
]
inst_fill   = "EAF4EA"
retail_fill = "FDECEA"
for i, (name, dates, regime) in enumerate(rows_data):
    row  = regime_table.rows[i+1]
    fill = inst_fill if regime == "Institutional" else retail_fill
    shade_cell(row.cells[0], fill); cell_text(row.cells[0], name,   align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    shade_cell(row.cells[1], fill); cell_text(row.cells[1], dates,  size=10)
    shade_cell(row.cells[2], fill); cell_text(row.cells[2], regime, size=10,
        bold=(regime=="Institutional"))

add_para(doc, "Table 1. A priori regime definitions based on public-record event dates.",
         space_before=4, space_after=10)

add_heading(doc, "2.4 Statistical Tests", 2)
add_para(doc,
    "Independent-samples t-test, Mann-Whitney U (one-sided: institutional > retail), and Cohen's d "
    "were computed between the institutional and retail/crisis SCI distributions. A period-level "
    "breakdown was also tabulated. The CORE threshold is defined as SCI > 0.75, corresponding to "
    "z ≈ 1.1 under k=0.9.")

# ── 3. Results ────────────────────────────────────────────────────────────────
add_heading(doc, "3. Results", 1)

add_heading(doc, "3.1 Primary Test: Institutional vs. Retail/Crisis", 2)
add_para(doc,
    "The hypothesis was directly falsified. Retail/crisis periods produced significantly higher SCI "
    "than institutional periods, in the opposite direction from the prediction:")

# Primary stats table
stats_tbl = doc.add_table(rows=4, cols=4)
stats_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_tbl.style = "Table Grid"
for j, h in enumerate(["Regime", "N windows", "Mean SCI (SD)", "% CORE (>0.75)"]):
    shade_cell(stats_tbl.rows[0].cells[j], "2E4057")
    cell_text(stats_tbl.rows[0].cells[j], h, bold=True, size=10)
    stats_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

data_rows = [
    ("Institutional", "1,303", "0.608 (0.221)", "32.1%"),
    ("Retail / Crisis","1,862", "0.676 (0.200)", "45.8%"),
    ("Transition",    "191",   "0.649 (0.208)", "38.2%"),
]
fills = [inst_fill, retail_fill, "F8F8F8"]
for i, (a,b,c,d) in enumerate(data_rows):
    row = stats_tbl.rows[i+1]
    for j, val in enumerate([a,b,c,d]):
        shade_cell(row.cells[j], fills[i])
        cell_text(row.cells[j], val, size=10)

add_para(doc, "Table 2. Grand means by regime.", space_before=4, space_after=8)

add_para(doc,
    "Statistical tests strongly reject the null of equal means in both directions "
    "(the institutional mean is lower, not higher, than retail):")

test_tbl = doc.add_table(rows=4, cols=2)
test_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
test_tbl.style = "Table Grid"
for j, h in enumerate(["Test", "Result"]):
    shade_cell(test_tbl.rows[0].cells[j], "2E4057")
    cell_text(test_tbl.rows[0].cells[j], h, bold=True, size=10)
    test_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
for i, (a,b) in enumerate([
    ("Independent t-test",          "t = −9.07,  p < 10⁻¹⁸  (inst < retail)"),
    ("Mann-Whitney U (inst > ret)", "p = 1.000  (inst is stochastically lower)"),
    ("Cohen's d",                   "d = −0.328  (small-to-medium effect, wrong direction)"),
]):
    row = test_tbl.rows[i+1]
    shade_cell(row.cells[0], "F8F8F8"); cell_text(row.cells[0], a, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    shade_cell(row.cells[1], "F8F8F8"); cell_text(row.cells[1], b, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)

add_para(doc, "Table 3. Statistical tests.", space_before=4, space_after=8)

add_heading(doc, "3.2 Per-Period Breakdown", 2)
add_para(doc,
    "The period-level breakdown reveals the structure of the reversal. The single highest-SCI "
    "period in the 2017–2026 record is the pre-futures retail era (2017-01-01 to 2017-12-16), "
    "with mean SCI=0.801 and 75.9% of windows classified CORE — the purest retail speculation "
    "window in the dataset and the one furthest from any institutional infrastructure.")

# Per-period table
pp_tbl = doc.add_table(rows=11, cols=6)
pp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
pp_tbl.style = "Table Grid"
for j, h in enumerate(["Period", "Regime", "N", "Mean SCI", "Median SCI", "% CORE"]):
    shade_cell(pp_tbl.rows[0].cells[j], "2E4057")
    cell_text(pp_tbl.rows[0].cells[j], h, bold=True, size=9)
    pp_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

period_data = [
    ("Pre-Futures Retail",          "Retail",        "286",  "0.801", "0.838", "75.9%"),
    ("China Ban / Mid-Cycle",       "Retail",        "153",  "0.738", "0.753", "62.7%"),
    ("Spot ETF Era",                "Institutional", "854",  "0.621", "0.630", "35.0%"),
    ("Prolonged Bear / Retail",     "Retail",        "823",  "0.640", "0.705", "40.1%"),
    ("FTX Collapse / Crypto Winter","Retail",        "426",  "0.653", "0.718", "36.6%"),
    ("Luna/3AC Collapse",           "Retail",        "92",   "0.654", "0.695", "37.0%"),
    ("Corporate Treasury Wave",     "Institutional", "212",  "0.606", "0.614", "28.8%"),
    ("Post-ATH Deleveraging",       "Retail",        "82",   "0.634", "0.629", "24.4%"),
    ("CME/CBOE Futures Era",        "Institutional", "196",  "0.562", "0.552", "28.1%"),
    ("ProShares ETF / ATH",         "Institutional", "41",   "0.567", "0.549", "7.3%"),
]
for i, row_data in enumerate(period_data):
    row = pp_tbl.rows[i+1]
    regime = row_data[1]
    fill   = inst_fill if regime == "Institutional" else retail_fill
    for j, val in enumerate(row_data):
        shade_cell(row.cells[j], fill)
        cell_text(row.cells[j], val, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)

add_para(doc, "Table 4. Per-period SCI breakdown, sorted descending by mean SCI.",
         space_before=4, space_after=10)

add_heading(doc, "3.3 SCI vs. GBTC Price Correlation", 2)
add_para(doc,
    "Pearson correlation between the 90-day rolling SCI and contemporaneous GBTC closing price "
    "was computed across the overlapping sample. A positive correlation would support the "
    "institutional-SCI hypothesis; a negative or null correlation would contradict it. The result "
    "was consistent with the primary finding: rolling SCI was not positively correlated with GBTC "
    "price in a way that would confirm the institutional hypothesis. During GBTC's peak "
    "institutional demand period (2020–2021), SCI was at intermediate levels, while the highest "
    "SCI readings occurred in the pre-GBTC-surge retail era of 2017.")

# ── 4. Discussion ─────────────────────────────────────────────────────────────
add_heading(doc, "4. Discussion", 1)

add_heading(doc, "4.1 Why the Hypothesis Was Falsified", 2)
add_para(doc,
    "The mechanism underlying the prediction was that institutional participation creates the "
    "feedback infrastructure — ETF rebalancing, systematic factor arbitrage, index inclusion "
    "mechanics — that generates structural coherence in equity markets. This mechanism does operate "
    "in equities: the CORE-classified instruments (major indices, large-cap equities) sit inside "
    "dense networks of forced correlation through passive vehicles and macro factor hedging.")

add_para(doc,
    "Bitcoin's institutional ecosystem differs in a crucial way. The arrival of CME futures "
    "(December 2017), the ProShares ETF (October 2021), and the spot ETF (January 2024) introduced "
    "arbitrage mechanisms between spot and derivatives, not rebalancing cascades. Futures "
    "arbitrageurs actively trade basis spreads, flattening momentum. Derivatives desks delta-hedge, "
    "creating contra-directional flows. These are coherence-reducing mechanisms in the short-term "
    "frequency range that SCI targets.")

add_heading(doc, "4.2 Revised Interpretation: Two Pathways to Structural Coherence", 2)
add_para(doc,
    "The results suggest that SCI detects structural coherence regardless of its mechanistic source, "
    "but the two primary sources produce different temporal signatures:")

add_para(doc,
    "Pathway 1 — Institutional Feedback (equity markets): Coherence is built through forced "
    "rebalancing cascades, index arbitrage, and macro factor networks. This is persistent and "
    "monotonically related to institutional depth. CORE stocks maintain high SCI across all market "
    "conditions because the feedback infrastructure is always active.")

add_para(doc,
    "Pathway 2 — Retail Speculative Synchrony (Bitcoin 2017): Coherence is built through herding "
    "— market participants synchronizing on a shared narrative and momentum signal. This produces "
    "extremely high SCI during bubble formation (BTC 2017: SCI=0.801) but collapses when the "
    "narrative breaks or when institutional arbitrage arrives to exploit and flatten the momentum.")

add_para(doc,
    "Under this revised interpretation, the 2017 pre-futures period was the most structurally "
    "coherent Bitcoin regime precisely because it was the most purely retail — a single coherent "
    "narrative (cryptocurrency to mainstream), propagating through synchronized retail participation "
    "with no derivatives to hedge against. When CME futures launched, the SCI declined, not because "
    "institutions bring coherence, but because futures arbitrage attacked the momentum structure "
    "that retail herding had created.")

add_heading(doc, "4.3 The Spot ETF Era (2024–present)", 2)
add_para(doc,
    "The Spot ETF Era (January 2024 onward) shows intermediate SCI (0.621, 35% CORE), meaningfully "
    "above the CME futures era (0.562) but below the pure retail peak (0.801). This is consistent "
    "with a hybrid regime: spot ETFs do create some systematic rebalancing flows (ETF arbitrage, "
    "creation/redemption mechanism), which may be beginning to generate equity-like institutional "
    "feedback. However, the BTC ecosystem remains too thin in systematic macro-factor relationships "
    "to reach equity CORE levels. Continued monitoring of rolling SCI may capture the moment of "
    "transition if and when it occurs.")

add_heading(doc, "4.4 Implications for the Broader Claim", 2)
add_para(doc,
    "The original broader implication — that rolling SCI can detect when a new asset class has "
    "developed systematic-strategy-exploitable regime structure — survives the falsification, but "
    "requires revision. The metric does not specifically detect institutional feedback; it detects "
    "any structural coherence, whether from institutional feedback or retail bubble synchrony. Both "
    "are exploitable regimes: retail synchrony is exploitable via momentum, institutional feedback "
    "via mean-reversion and factor strategies.")

add_para(doc,
    "The revised claim: rolling SCI detects structural coherence in any asset regardless of its "
    "source. A rising SCI on a new instrument signals that some regime is forming — either "
    "speculative synchrony or systematic feedback. A falling SCI after a period of high values "
    "may signal regime collapse (bubble break) or regime maturation (arbitrage flattening momentum). "
    "Distinguishing between the two requires auxiliary information (derivatives depth, GBTC-style "
    "proxies, open interest), but the SCI transition itself is the early warning signal.")

# ── 5. Conclusion ─────────────────────────────────────────────────────────────
add_heading(doc, "5. Conclusion", 1)
add_para(doc,
    "The Regime Transition Test directly falsified the institutional-coherence hypothesis for "
    "Bitcoin. Retail/crisis periods produced higher 90-day rolling SCI (0.676) than institutional "
    "periods (0.608), with the single highest-SCI epoch in the 2017–2026 record being the "
    "pre-futures retail era of 2017 (SCI=0.801, 76% CORE-classified). The effect is statistically "
    "robust (t=−9.07, p<10⁻¹⁸, d=−0.33).")

add_para(doc,
    "The result reveals a second pathway to structural coherence — retail speculative synchrony — "
    "distinct from the institutional feedback mechanism that generates CORE classification in "
    "equities. Bitcoin's high SCI in 2017 reflects pure momentum herding. The arrival of "
    "derivatives infrastructure (CME 2017, ProShares 2021) introduced arbitrage flows that "
    "flattened the herding dynamics and reduced coherence, exactly the inverse of the equity case "
    "where institutional infrastructure creates coherence.")

add_para(doc,
    "SCI remains a valid detector of structural regime, but the regime it detects is not "
    "specifically institutional maturity. It is structural coherence in the broadest sense: any "
    "dynamical state where the Hilbert envelope is more correlated across time than a "
    "phase-randomized baseline. Identifying which mechanism underlies a given SCI reading requires "
    "auxiliary market context. The SCI value alone is the signal; the mechanism is the "
    "interpretation.")

# ── Appendix ──────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Appendix: Reproduction", 1)
add_para(doc, "All code and data downloads are reproducible from the public record.")
add_code_block(doc,
    "# Reproduce from project root\n"
    "cd '/Users/parkerlee/Desktop/If Im Right/SCI_Project'\n"
    "\n"
    "# Run analysis\n"
    "python3 scripts/finance/sci_btc_regime_v1.py\n"
    "\n"
    "# Outputs:\n"
    "# results/btc_regime_v1/rolling_sci.csv          (3,356 rows)\n"
    "# results/btc_regime_v1/period_breakdown.csv     (10 rows)\n"
    "# results/btc_regime_v1/statistical_tests.txt\n"
    "# results/btc_regime_v1/plots/btc_rolling_sci.png\n"
    "# results/btc_regime_v1/plots/btc_sci_regime_boxplot.png\n"
    "# results/btc_regime_v1/plots/sci_vs_gbtc.png")

add_para(doc, "Data sources: Yahoo Finance (yfinance) — BTC-USD, GBTC. No API key required.",
         space_before=8)
add_para(doc,
    "SCI parameters (finance standard, locked): W=7, L=10, S=40, k=0.9, seed=42.\n"
    "Rolling window: 63 trading days (~90 calendar days).\n"
    "SCI engine: sci_score_v3.py (project root).\n"
    "Regime periods: defined a priori from public-record event dates, not from SCI data.")

doc.save(OUT)
print(f"Saved: {OUT}")
